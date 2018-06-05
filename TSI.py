import os
import re
import time
from docx import Document

document = Document("G:\Enrollment Management Center\Graduation\default.docx")
fileLocation = input('Drop file here (Delete quotes around file name): \n')
os.chdir("G:\Enrollment Management Center\EDI transcripts\Clean EDI")


class CleanLine:
    '''This class cleans the extra blank spaces and changes the line according
    to the data inside'''

    def __init__(self, line):
        '''Removes all the extra blank spaces from everyline of the EDI'''
        self.line = re.sub(r'(\s+)', ' ', line.strip())

    def drop_count(self):
        foo = re.compile(r'(DROPCOUNT)(=)\s*(\d+)', re.IGNORECASE)
        mo = foo.search(self.line)
        Drops, Equals, Count = mo.groups()
        self.line = re.sub(r'.+', '{} {} {}'.format(Drops, Equals, Count),
                           self.line)
        return self.line

    def course(self):
        '''This function removes all of the unnecessary information from the
        lines containing coursework'''
        self.line = re.sub(r'(\W)$', '', self.line)
        self.line = re.sub(r'((\s)(\d))$', '', self.line)
        self.line = re.sub(r'((\s)([E|I]))$', '', self.line)
        self.line = re.sub(r'((I/F))$', 'F', self.line)
        self.line = re.sub(r'((WF|WQ|FX))$', 'F', self.line)
        self.line = re.sub(r'((WL|WP))$', 'W', self.line)
        self.line = re.sub(r'^(WBCT).+', '', self.line)
        self.line = re.sub(r'(ZZZ)$', 'IP', self.line)
        return self.line


cclmList = ['MATH 1314', 'MATH 1316', 'MATH 1324', 'MATH 1325', 'MATH 1332',
            'MATH 1333', 'MATH 1342', 'MATH 1350', 'MATH 1351', 'MATH 2318',
            'MATH 2320', 'MATH 2412', 'MATH 2413', 'MATH 2414', 'MATH 2415',
            'MATH 0308', 'MATH 0312', 'M 305G', 'M 408N', 'MATH 1369',
            'MATH 1320', 'MATH 1233', 'MATH 1433']

cclrList = ['ANTH 2351', 'ARTS 1303', 'ARTS 1304', 'CHEM 1411', 'CHEM 1412',
            'DANC 2303', 'ECON 2301', 'ECON 2302', 'ENGL 1301', 'ENGL 1301',
            'ENGL 2322', 'ENGL 2323', 'ENGL 2327', 'ENGL 2328', 'ENGL 2332',
            'ENGL 2333', 'ENGL 2341', 'GEOL 1447', 'GOVT 2301', 'GOVT 2302',
            'HIST 1301', 'HIST 1302', 'HIST 2301', 'HIST 2311', 'HIST 2312',
            'HIST 2321', 'HIST 2322', 'PHIL 1301', 'PHIL 2306', 'PHYS 1401',
            'PHYS 1411', 'PHYS 1412', 'PHYS 2425', 'PSYC 2301', 'PSYC 2306',
            'PSYC 2308', 'PSYC 2314', 'PSYC 2315', 'PSYC 2317', 'PSYC 2319',
            'SOCI 2319', 'HIST 1013', 'PSY 301', 'CH 301', 'GOVT 2305',
            'GOVT 2306', 'POLS 2305', 'POLS 2306', 'POLS 1333', 'POLS 1433',
            'PSYC 1103', 'SOCL 1133']

cclwList = ['ENGL 1301', 'ENGL 1303', 'ENGL 1302', 'ENGL 1304', 'ENG 1013',
            'ENG 1023', 'ENGL 0310', 'RHE 306', 'WRIT 1301', 'ENGL 1113',
            'ENGL 1123']

testScoreList = ['9TX TSI', '* RAP: 9TX TSI', '9TX TASP']

with open(fileLocation, 'r') as transcript:
    data = transcript.readlines()
    for line in data:
        '''The loop will run through each line in the transcript and look for a
        REGEX match. If found it will then write the line to the document.'''

        new_line = CleanLine(line)

        # Writes the G# line down in BOLD
        if new_line.line.startswith('G0'):
            p = document.add_paragraph('=' * 50 + '\n')
            p.add_run('\n' + new_line.line + '\n').bold = True

        elif re.match(r'^((\d){6})', new_line.line):
            p.add_run('{}\n'.format(new_line.line)).bold = True

        elif re.search(r'DROPCOUNT', new_line.line, re.IGNORECASE):
            if re.search(r'(00)|(0\s)', new_line.line):
                continue
            else:
                p.add_run('\n{}\n'.format(new_line.line))

        elif new_line.line.startswith(tuple(testScoreList)):
            p.add_run('\n{}\n'.format(new_line.line)).italic = True

        elif new_line.line.startswith('AS: '):
            p.add_run('\n{}\n'.format(new_line.line))

        elif new_line.line.startswith('Student Name:'):
            p.add_run('\n{}\n'.format(new_line.line)).bold = True

        elif re.search(r'^([A-Z]{1,10})(\s)(\d)', new_line.course()):
            if new_line.course().startswith(tuple(cclmList)):
                if re.search(r'([A-D]|T[A-D]|TCR)$', new_line.course()):
                    p.add_run(new_line.course() + '\n')
            elif new_line.course().startswith(tuple(cclrList)):
                if re.search(r'([A-D]|T[A-D]|TCR)$', new_line.course()):
                    p.add_run(new_line.course() + '\n')
            elif new_line.course().startswith(tuple(cclwList)):
                if re.search(r'([A-D]|T[A-D]|TCR)$', new_line.course()):
                    p.add_run(new_line.course() + '\n')

        elif new_line.line.startswith('Academic Degree'):
            p.add_run('\n{}\n'.format(new_line.line))

        elif new_line.line.startswith('Degree Awarded Date'):
            p.add_run('\n{}\n'.format(new_line.line))

        else:
            continue

fileName = input('\nWhat would you like to name the new file?:\n')
document.save(fileName + '.docx')

# For Peace of Mind
print("\nMaking Document...")
time.sleep(3)

exit = input('\nPress ENTER to close the program')
