import os
import re
import time
from docx import Document

document = Document("G:\Enrollment Management Center\Graduation\default.docx")
fileLocation = input('Drop file here (Delete quotes around file name): \n')
os.chdir("G:\Enrollment Management Center\EDI transcripts\Clean EDI")

with open(fileLocation, 'r') as transcript:
    data = transcript.readlines()
    for line in data:
        '''
        The loop will run through each line in the transcript and look for a
        REGEX match. If found it will then write the line to the document.
        '''

        # Writes the G# line down in BOLD
        if line.startswith('Sending Inst:'):
            p = document.add_paragraph('=' * 50 + '\n')
            p.add_run('\n' + line + '\n').bold = True

        # Matches SSN
        elif re.match(r'\d\d\d-\d\d-\d\d\d\d', line):
            p.add_run('{}\n'.format(line)).bold = True

        # Lists Grad Type
        elif line.startswith('Note: '):
            p.add_run(line)

        # List the level of graduation
        elif re.match(r'(\s+)(#CA:)', line):
            p.add_run(line)

        # Lists graduation date
        elif re.match(r'(\s+)(HS S)', line):
            p.add_run(line)
            p.add_run('-' * 50 + '\n')

        else:
            continue

fileName = input('\nWhat would you like to name the new file?:\n')
document.save(fileName + '.docx')

# For Peace of Mind
print("\nMaking Document...")
time.sleep(3)

exit = input('\nPress ENTER to close the program')
